"""
Ingest manifestos supplied directly to Jerseyvotes.org (rather than scraped
from vote.je or found on the open web). Some candidates campaign only via
social media, hustings handouts, or by emailing a Word document — this
script processes whatever was deposited into

    pipeline/manual_manifestos/<vote-je-slug>/

and writes the combined extraction into `candidates.enhanced_manifesto_*`
with `source_label='manually_supplied'` (or another label configured in
that candidate's `metadata.yaml`).

Supported file types:
  - .pdf                              -> Claude PDF (extract_pdf_claude.py)
  - .docx                             -> python-docx
  - .png/.jpg/.jpeg/.gif/.webp        -> Claude vision
  - .txt/.md                          -> plain UTF-8 read

The combined text is cached to `extracted.txt` next to the source files so
re-runs are free; pass --no-cache to force re-extraction. Source files AND
the cache are committed to git so the public can audit what we used.

The DB CHECK constraint enhanced_manifesto_text_must_have_source_url is
satisfied by using either the candidate-supplied `original_url` from
metadata.yaml (when the document is also publicly available, e.g. a
Facebook post) or the GitHub URL of the candidate's folder in this repo.

Run:
  python pipeline/ingest_manual_manifestos.py                 # all candidates
  python pipeline/ingest_manual_manifestos.py --slug jane-doe # one candidate
  python pipeline/ingest_manual_manifestos.py --no-cache      # re-extract
"""

import argparse
import base64
import os
import sys
from pathlib import Path

import anthropic
import psycopg2
import yaml
from dotenv import load_dotenv

# Local import — both scripts live in pipeline/ and there's no __init__.py
sys.path.insert(0, str(Path(__file__).resolve().parent))
from extract_pdf_claude import extract_pdf  # noqa: E402

load_dotenv(override=True)


MANUAL_DIR = Path(__file__).resolve().parent / 'manual_manifestos'

# Repo URL used as the public source URL when no `original_url` is supplied —
# anyone can browse this folder on GitHub to see what was processed.
REPO_HTTP_BASE = 'https://github.com/gusfraser/jerseyvotes.org/tree/main/pipeline/manual_manifestos'

# Labels we recognise as "supplied to us" rather than "found on the open web".
# Must be kept in sync with ENHANCED_SOURCE_LABELS in
# web/src/app/candidates/[slug]/page.tsx and with the exclusion list in
# find_enhanced_manifestos.py (so the auto-finder doesn't overwrite them).
MANUAL_SOURCE_LABELS = {
    'manually_supplied',
    'social_media_capture',
    'printed_handout',
}

SKIP_NAMES = {'metadata.yaml', 'extracted.txt', 'README.md', '.DS_Store'}

IMAGE_MEDIA_TYPES = {
    '.png': 'image/png',
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.gif': 'image/gif',
    '.webp': 'image/webp',
}
TEXT_EXTS = {'.txt', '.md'}
DOCX_EXT = '.docx'
PDF_EXT = '.pdf'

MODEL_VISION = 'claude-sonnet-4-5'
VISION_MAX_TOKENS = 16000

VISION_PROMPT = """You are extracting the full text of a political manifesto from an image (scanned page, photo of printed material, or social-media screenshot) for downstream LLM analysis.

Produce a faithful, plain-text transcription of the substantive content in correct reading order.

Rules:
1. Transcribe verbatim — do not summarise, paraphrase, or interpret.
2. Preserve paragraph breaks (single blank line between paragraphs).
3. For multi-column layouts, read each column top-to-bottom, then move to the next column.
4. Preserve bullet lists using "- " prefix.
5. Headings on their own line, using markdown headers (#, ##, ###) by depth.
6. Skip site chrome, like/share counts, navigation, watermarks, and other non-content.
7. Do not add commentary or notes — output the manifesto text only.
8. Use UK English as written.

If the image is illegible or contains no political content, output the literal token EMPTY on its own line."""


def db_connect():
    return psycopg2.connect(
        os.environ['DATABASE_URL'],
        keepalives=1, keepalives_idle=30, keepalives_interval=10,
        keepalives_count=5, connect_timeout=15,
    )


def extract_image(path: Path, client: anthropic.Anthropic) -> str:
    media_type = IMAGE_MEDIA_TYPES[path.suffix.lower()]
    data = base64.standard_b64encode(path.read_bytes()).decode('utf-8')
    resp = client.messages.create(
        model=MODEL_VISION,
        max_tokens=VISION_MAX_TOKENS,
        messages=[{
            'role': 'user',
            'content': [
                {'type': 'image', 'source': {
                    'type': 'base64', 'media_type': media_type, 'data': data,
                }},
                {'type': 'text', 'text': VISION_PROMPT},
            ],
        }],
    )
    text = ''.join(b.text for b in resp.content if getattr(b, 'type', None) == 'text').strip()
    return '' if text == 'EMPTY' else text


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit('python-docx is required for .docx files: pip install python-docx')
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # Tables: flatten cell text row by row. Manifestos rarely use tables but
    # candidate-supplied policy grids do show up — better to capture than drop.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n\n'.join(parts)


def extract_text_file(path: Path) -> str:
    return path.read_text(encoding='utf-8', errors='replace').strip()


def extract_one(path: Path, client: anthropic.Anthropic) -> str:
    ext = path.suffix.lower()
    if ext == PDF_EXT:
        return extract_pdf(str(path)).strip()
    if ext == DOCX_EXT:
        return extract_docx(path).strip()
    if ext in IMAGE_MEDIA_TYPES:
        return extract_image(path, client).strip()
    if ext in TEXT_EXTS:
        return extract_text_file(path)
    raise ValueError(f'unsupported file type: {path.name}')


def load_metadata(folder: Path) -> dict:
    meta_path = folder / 'metadata.yaml'
    if not meta_path.exists():
        raise FileNotFoundError(f'{folder.name}: missing metadata.yaml')
    data = yaml.safe_load(meta_path.read_text()) or {}
    if not isinstance(data, dict):
        raise ValueError(f'{folder.name}: metadata.yaml must be a YAML mapping')
    slug = (data.get('candidate_slug') or '').strip()
    if not slug:
        raise ValueError(f'{folder.name}: metadata.yaml needs candidate_slug')
    source_label = (data.get('source_label') or 'manually_supplied').strip()
    if source_label not in MANUAL_SOURCE_LABELS:
        raise ValueError(
            f'{folder.name}: source_label {source_label!r} not in '
            f'{sorted(MANUAL_SOURCE_LABELS)} — add it to MANUAL_SOURCE_LABELS '
            'in ingest_manual_manifestos.py and ENHANCED_SOURCE_LABELS in '
            'web/src/app/candidates/[slug]/page.tsx first'
        )
    note = (data.get('note') or '').strip()
    original_url = (data.get('original_url') or '').strip() or None
    explicit_files = data.get('files')
    if explicit_files is not None and not isinstance(explicit_files, list):
        raise ValueError(f'{folder.name}: `files:` must be a list if present')
    return {
        'candidate_slug': slug,
        'source_label': source_label,
        'note': note,
        'original_url': original_url,
        'files': explicit_files,
    }


def collect_files(folder: Path, explicit: list[str] | None) -> list[Path]:
    if explicit:
        return [folder / name for name in explicit]
    return sorted(
        p for p in folder.iterdir()
        if p.is_file() and p.name not in SKIP_NAMES and not p.name.startswith('.')
    )


def combine_extractions(parts: list[tuple[str, str]]) -> str:
    """Concatenate per-file extractions with file-named headings so the
    classifier sees clear source boundaries and a reviewer can tell where
    each excerpt came from."""
    blocks = []
    for fname, text in parts:
        if not text:
            continue
        blocks.append(f'## {fname}\n\n{text}')
    return '\n\n'.join(blocks).strip()


def upsert_enhanced(conn, cur, slug: str, text: str, *,
                    source_url: str, source_label: str, notes: str) -> int:
    word_count = len(text.split())
    cur.execute(
        '''
        UPDATE candidates SET
            enhanced_manifesto_text = %(text)s,
            enhanced_manifesto_word_count = %(wc)s,
            enhanced_manifesto_source_url = %(url)s,
            enhanced_manifesto_source_label = %(label)s,
            enhanced_manifesto_status = 'found',
            enhanced_manifesto_notes = %(notes)s,
            enhanced_manifesto_fetched_at = NOW()
        WHERE vote_je_slug = %(slug)s
        ''',
        {
            'text': text, 'wc': word_count,
            'url': source_url, 'label': source_label,
            'notes': notes[:1000], 'slug': slug,
        },
    )
    if cur.rowcount == 0:
        raise ValueError(f'no candidate with vote_je_slug = {slug!r}')
    conn.commit()
    return word_count


def main():
    parser = argparse.ArgumentParser(
        description='Ingest manually-supplied manifestos into candidates.enhanced_manifesto_*'
    )
    parser.add_argument('--slug', help='Process only this candidate folder/slug')
    parser.add_argument('--limit', type=int, default=0,
                        help='Process at most N candidates (after slug filter)')
    parser.add_argument('--no-cache', action='store_true',
                        help='Re-extract even if extracted.txt is present')
    parser.add_argument('--dry-run', action='store_true',
                        help='Extract and cache but do not write to the DB')
    args = parser.parse_args()

    if not MANUAL_DIR.exists():
        sys.exit(f'No manual manifestos directory at {MANUAL_DIR}')

    folders = sorted(
        p for p in MANUAL_DIR.iterdir()
        if p.is_dir() and not p.name.startswith('_') and not p.name.startswith('.')
    )
    if args.slug:
        folders = [p for p in folders if p.name == args.slug]
        if not folders:
            sys.exit(f'No folder named {args.slug!r} under {MANUAL_DIR}')
    if args.limit:
        folders = folders[: args.limit]

    if not folders:
        print(f'No candidate folders to process under {MANUAL_DIR}')
        return

    client = anthropic.Anthropic()
    conn = None
    cur = None
    if not args.dry_run:
        conn = db_connect()
        cur = conn.cursor()

    ok, errored = 0, 0
    try:
        for folder in folders:
            print(f'[{folder.name}]')
            try:
                meta = load_metadata(folder)
            except Exception as e:
                print(f'  metadata error: {e}')
                errored += 1
                continue

            slug = meta['candidate_slug']
            if slug != folder.name:
                print(f'  warning: candidate_slug ({slug!r}) != folder name ({folder.name!r})')

            files = collect_files(folder, meta['files'])
            cache = folder / 'extracted.txt'

            if cache.exists() and not args.no_cache:
                combined = cache.read_text(encoding='utf-8').strip()
                print(f'  using cached extracted.txt ({len(combined.split())} words)')
            else:
                if not files:
                    print(f'  no source files; skipping')
                    errored += 1
                    continue
                parts: list[tuple[str, str]] = []
                for f in files:
                    if not f.exists():
                        print(f'  missing file: {f.name}')
                        continue
                    print(f'  extracting {f.name} ({f.suffix.lstrip(".") or "?"})')
                    try:
                        text = extract_one(f, client)
                    except Exception as e:
                        print(f'    extract error: {e}')
                        continue
                    if text:
                        parts.append((f.name, text))
                if not parts:
                    print(f'  nothing extracted; skipping')
                    errored += 1
                    continue
                combined = combine_extractions(parts)
                cache.write_text(combined + '\n', encoding='utf-8')
                print(f'  cached extraction -> {cache.relative_to(MANUAL_DIR.parent.parent)}')

            wc = len(combined.split())
            if wc < 30:
                print(f'  combined too short ({wc} words); skipping DB write')
                errored += 1
                continue

            # Prefer the candidate's original public URL when supplied; fall
            # back to the GitHub folder URL so the CHECK constraint requiring
            # a traceable source is always satisfied.
            source_url = meta['original_url'] or f'{REPO_HTTP_BASE}/{folder.name}'

            file_list = ', '.join(p.name for p in files)
            note_bits = []
            if meta['note']:
                note_bits.append(meta['note'])
            note_bits.append(f'Files: {file_list}')
            if meta['original_url']:
                note_bits.append(f'Public source: {meta["original_url"]}')
            note_bits.append(f'Repo audit: {REPO_HTTP_BASE}/{folder.name}')

            if args.dry_run:
                print(f'  [dry-run] would upsert {wc} words, source_label={meta["source_label"]!r}, source_url={source_url!r}')
                ok += 1
                continue

            try:
                upsert_enhanced(
                    conn, cur, slug, combined,
                    source_url=source_url,
                    source_label=meta['source_label'],
                    notes=' | '.join(note_bits),
                )
                print(f'  updated candidates row ({wc} words) -> {source_url}')
                ok += 1
            except Exception as e:
                print(f'  DB error: {e}')
                errored += 1
    finally:
        if cur is not None:
            cur.close()
        if conn is not None:
            conn.close()

    print(f'\nDone. ok={ok}, errors={errored}')


if __name__ == '__main__':
    main()
